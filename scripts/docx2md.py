"""docx → md 转换（保段落顺序 + 表格 pipe 化,纯 python-docx 实现,不依赖 pandoc）.

用法：python scripts/docx2md.py <input.docx> <output.md>

- Title / Heading N 样式映射到 # / ## / ###
- 表格转 markdown pipe 表,单元格内多段落用 <br> 合并
- 列表样式降级为 "- " 前缀项
- 遍历 doc.element.body 保段落/表格原顺序
"""
import sys
import re
from pathlib import Path
from docx import Document
from docx.oxml.ns import qn


def _cell_text(cell):
    # 单元格内多段落合并，换行替换为 <br>
    parts = [p.text.strip() for p in cell.paragraphs if p.text.strip()]
    text = "<br>".join(parts)
    # 管道字符转义，防止破坏 md 表格
    return text.replace("|", "\\|").replace("\n", "<br>")


def _table_to_md(table) -> str:
    rows = []
    for row in table.rows:
        rows.append([_cell_text(c) for c in row.cells])
    if not rows:
        return ""
    # 去重（python-docx 对合并单元格返回重复对象）
    ncols = max(len(r) for r in rows)
    rows = [r + [""] * (ncols - len(r)) for r in rows]

    header = rows[0]
    sep = ["---"] * ncols
    body = rows[1:] if len(rows) > 1 else []

    out = ["| " + " | ".join(header) + " |", "| " + " | ".join(sep) + " |"]
    for r in body:
        out.append("| " + " | ".join(r) + " |")
    return "\n".join(out)


def _para_to_md(para) -> str:
    text = para.text.strip()
    if not text:
        return ""
    style_name = (para.style.name if para.style else "") or ""
    sn = style_name.lower()

    if sn.startswith("title"):
        return f"# {text}"
    if sn.startswith("heading"):
        # Heading 1 / Heading 2 / ...
        m = re.search(r"(\d+)", sn)
        level = int(m.group(1)) if m else 2
        level = max(1, min(level + 1, 6))  # Heading 1 → ## 让 Title 独占 #
        return f"{'#' * level} {text}"
    # 列表样式
    if "list bullet" in sn or "list" in sn:
        return f"- {text}"
    return text


def convert(src: Path, dst: Path):
    doc = Document(str(src))
    body = doc.element.body

    # 建 index 映射 python-docx 对象
    para_map = {p._element: p for p in doc.paragraphs}
    tbl_map = {t._element: t for t in doc.tables}

    out_lines: list[str] = []
    for child in body.iterchildren():
        if child.tag == qn("w:p"):
            p = para_map.get(child)
            if p is None:
                continue
            md = _para_to_md(p)
            if md:
                out_lines.append(md)
        elif child.tag == qn("w:tbl"):
            t = tbl_map.get(child)
            if t is None:
                continue
            tbl_md = _table_to_md(t)
            if tbl_md:
                out_lines.append("")
                out_lines.append(tbl_md)
                out_lines.append("")

    dst.parent.mkdir(parents=True, exist_ok=True)
    dst.write_text("\n\n".join(out_lines) + "\n", encoding="utf-8")
    print(f"wrote {dst} ({dst.stat().st_size} bytes)")


if __name__ == "__main__":
    convert(Path(sys.argv[1]), Path(sys.argv[2]))
